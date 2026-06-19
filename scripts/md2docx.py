"""Convert MiniCode markdown doc to DOCX (save as PDF from Word).

Usage:
    python scripts/md2docx.py "C:/Users/123/Desktop/MiniCode项目深度解析.md"
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert(md_path: str, docx_path: str) -> None:
    md_text = Path(md_path).read_text(encoding="utf-8")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.5

    in_code_block = False

    for line in md_text.split("\n"):
        line = line.rstrip()

        # Code block toggle
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        # H1
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            run = p.add_run(line[2:])
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(30, 64, 175)
            continue

        # H2
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            run = p.add_run(line[3:])
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(30, 64, 175)
            continue

        # H3
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 3"]
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            continue

        # Horizontal rule → section break
        if line.strip() == "---":
            doc.add_paragraph("─" * 50)
            continue

        # Table row — skip or render simply
        if line.startswith("|"):
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.color.rgb = RGBColor(100, 116, 139)
            run.font.italic = True
            p.paragraph_format.left_indent = Cm(1)
            continue

        # Bullet
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(_clean_inline(line[2:]))
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(_clean_inline(m.group(2)))
            continue

        # Bold-only lines
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            continue

        # Empty
        if not line.strip():
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, _clean_inline(line))

    doc.save(docx_path)
    print(f"DOCX saved to: {docx_path}")


def _clean_inline(text: str) -> str:
    """Strip inline markdown without losing content."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to paragraph, handling inline code markers."""
    paragraph.add_run(text)


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "MiniCode项目深度解析.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else src.replace(".md", ".docx")
    convert(src, dst)
